"""
Script untuk membuat file PRESENTASI_GAMEFINDER.docx
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ============================================================
# STYLE SETUP
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

for i in range(1, 4):
    heading_style = doc.styles[f'Heading {i}']
    heading_style.font.name = 'Times New Roman'
    heading_style.font.color.rgb = RGBColor(0, 0, 0)

# Set margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.54)

# ============================================================
# HELPER FUNCTIONS
# ============================================================
def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Shading
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D9E2F3')
        shading.set(qn('w:val'), 'clear')
        cell.paragraphs[0].runs[0].element.getparent().getparent().find(qn('w:tcPr')).append(shading) if cell._element.find(qn('w:tcPr')) is not None else None
    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_formula(doc, formula_text, label=None):
    """Add a formula as an indented, italic paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(formula_text)
    run.font.name = 'Cambria Math'
    run.font.size = Pt(12)
    run.italic = True
    return p

def add_bullet(doc, text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.add_run(text).font.name = 'Times New Roman'
    else:
        p.text = ''
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

def add_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ============================================================
# COVER PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('MATERI PRESENTASI')
run.bold = True
run.font.size = Pt(16)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('GameFinder')
run.bold = True
run.font.size = Pt(20)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0, 51, 153)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Sistem Rekomendasi Game Steam\nMenggunakan Metode Fuzzy Mamdani')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Mata Kuliah: Kecerdasan Buatan')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Semester 4 — Tahun Akademik 2025/2026')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_page_break()

# ============================================================
# DAFTAR ISI
# ============================================================
doc.add_heading('DAFTAR ISI', level=1)
add_line(doc)

toc_items = [
    ('1.', 'Permasalahan yang Diangkat (Latar Belakang)'),
    ('2.', 'Permasalahan yang Diselesaikan oleh Sistem'),
    ('3.', 'Teori Metode Fuzzy Mamdani'),
    ('4.', 'Cara Kerja Codingan (Arsitektur Sistem)'),
    ('5.', 'Contoh Input → Proses → Output'),
    ('6.', 'Referensi Jurnal'),
    ('7.', 'Batasan / Kekurangan Sistem'),
    ('8.', 'Ringkasan untuk Penutup Presentasi'),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{num}  {title}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

doc.add_page_break()

# ============================================================
# BAB 1: PERMASALAHAN
# ============================================================
doc.add_heading('1. PERMASALAHAN YANG DIANGKAT (Latar Belakang)', level=1)
add_line(doc)

doc.add_heading('1.1 Masalah Utama', level=2)
p = doc.add_paragraph()
run = p.add_run('Platform Steam memiliki ')
run.font.name = 'Times New Roman'
run = p.add_run('lebih dari 70.000+ judul game')
run.bold = True
run.font.name = 'Times New Roman'
run = p.add_run(' yang tersedia. Pengguna (terutama gamer pemula) sering mengalami ')
run.font.name = 'Times New Roman'
run = p.add_run('information overload')
run.bold = True
run.italic = True
run.font.name = 'Times New Roman'
run = p.add_run(' — kebingungan dalam memilih game yang tepat karena banyaknya pilihan.')
run.font.name = 'Times New Roman'

doc.add_heading('1.2 Tantangan Spesifik yang Dihadapi Pengguna', level=2)
add_table(doc,
    ['No', 'Permasalahan', 'Penjelasan'],
    [
        ['1', 'Keterbatasan Anggaran', 'Mahasiswa/pelajar memiliki budget terbatas, tidak bisa asal beli game mahal'],
        ['2', 'Keterbatasan Spesifikasi PC', 'Banyak game yang tidak bisa dijalankan di PC/Laptop spesifikasi rendah'],
        ['3', 'Rating yang Membingungkan', 'Rating game di Steam bervariasi, pengguna sulit menilai mana yang benar-benar bagus'],
        ['4', 'Waktu Bermain Tidak Sesuai', 'Pengguna kasual tidak ingin game 100+ jam, sementara gamer hardcore butuh game panjang'],
        ['5', 'Multi-Kriteria', 'Manusia kesulitan mempertimbangkan 4–5 kriteria secara bersamaan dan objektif'],
    ],
    col_widths=[1.2, 4, 10]
)

doc.add_heading('1.3 Mengapa Masalah Ini Penting?', level=2)
add_bullet(doc, 'Membeli game yang salah = uang terbuang dan waktu terbuang')
add_bullet(doc, 'Tidak ada sistem di Steam yang secara cerdas mencocokkan profil pengguna (anggaran, spek PC, gaya bermain) dengan game secara bersamaan')
add_bullet(doc, 'Sistem rekomendasi bawaan Steam hanya berbasis riwayat pembelian dan popularitas, bukan berdasarkan analisis multi-kriteria')

doc.add_page_break()

# ============================================================
# BAB 2: SOLUSI
# ============================================================
doc.add_heading('2. PERMASALAHAN YANG DISELESAIKAN OLEH SISTEM', level=1)
add_line(doc)

doc.add_heading('2.1 Solusi yang Ditawarkan', level=2)
p = doc.add_paragraph()
run = p.add_run('GameFinder adalah ')
run.font.name = 'Times New Roman'
run = p.add_run('Sistem Pendukung Keputusan (Decision Support System)')
run.bold = True
run.font.name = 'Times New Roman'
run = p.add_run(' berbasis web yang menggunakan ')
run.font.name = 'Times New Roman'
run = p.add_run('Logika Fuzzy Mamdani')
run.bold = True
run.font.name = 'Times New Roman'
run = p.add_run(' untuk merekomendasikan game Steam terbaik berdasarkan profil pengguna.')
run.font.name = 'Times New Roman'

doc.add_heading('2.2 Apa yang Dilakukan Sistem', level=2)
add_table(doc,
    ['Fitur', 'Penjelasan'],
    [
        ['Input Multi-Kriteria', 'Pengguna memasukkan 5 kriteria sekaligus (Anggaran, Spek PC, Tipe Gamer, Rating, Waktu Bermain)'],
        ['Evaluasi Fuzzy', 'Setiap game dievaluasi menggunakan 243 aturan Fuzzy Mamdani'],
        ['Skor Kecocokan', 'Setiap game mendapat skor 0-100% yang menunjukkan tingkat kecocokan'],
        ['Explainable AI', 'Sistem memberikan alasan mengapa game direkomendasikan (bukan black-box)'],
        ['Top 10 Ranking', 'Menampilkan 10 game terbaik yang paling cocok untuk pengguna'],
    ],
    col_widths=[4, 12]
)

doc.add_heading('2.3 Keunggulan Dibandingkan Sistem Lain', level=2)
add_bullet(doc, ' Pengguna tahu mengapa game X direkomendasikan', bold_prefix='Transparan:')
add_bullet(doc, ' Langsung pakai tanpa registrasi', bold_prefix='Tidak perlu akun/login:')
add_bullet(doc, ' Tidak hanya berdasarkan 1 faktor, tapi 5 faktor bersamaan', bold_prefix='Multi-kriteria sekaligus:')
add_bullet(doc, ' Logika Fuzzy bisa menangani nilai "abu-abu" (misal: harga agak mahal)', bold_prefix='Toleransi ketidakpastian:')

doc.add_page_break()

# ============================================================
# BAB 3: TEORI MAMDANI
# ============================================================
doc.add_heading('3. TEORI METODE FUZZY MAMDANI', level=1)
add_line(doc)

doc.add_heading('3.1 Apa Itu Logika Fuzzy?', level=2)
p = doc.add_paragraph(
    'Logika Fuzzy adalah cabang dari kecerdasan buatan yang memungkinkan sebuah nilai '
    'memiliki derajat kebenaran antara 0 dan 1, bukan hanya 0 (Salah) atau 1 (Benar) '
    'seperti logika klasik (Boolean).'
)
p.runs[0].font.name = 'Times New Roman'

p = doc.add_paragraph()
run = p.add_run('Contoh Sederhana:')
run.bold = True
run.font.name = 'Times New Roman'
add_bullet(doc, 'Logika Klasik: Harga Rp 250.000 → Murah (Ya/Tidak?)')
add_bullet(doc, 'Logika Fuzzy: Harga Rp 250.000 → 83% Murah DAN 17% Sedang (bisa keduanya!)')

doc.add_heading('3.2 Apa Itu Metode Mamdani?', level=2)
p = doc.add_paragraph()
run = p.add_run('Metode Mamdani (Mamdani Inference System) dikembangkan oleh ')
run.font.name = 'Times New Roman'
run = p.add_run('Ebrahim H. Mamdani')
run.bold = True
run.font.name = 'Times New Roman'
run = p.add_run(' pada tahun ')
run.font.name = 'Times New Roman'
run = p.add_run('1975')
run.bold = True
run.font.name = 'Times New Roman'
run = p.add_run('. Metode ini merupakan salah satu metode inferensi fuzzy yang paling populer dan paling banyak digunakan.')
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
run = p.add_run('Ciri khas utama metode Mamdani:')
run.font.name = 'Times New Roman'
add_bullet(doc, 'Menggunakan aturan IF-THEN berbasis linguistik (bahasa manusia)')
add_bullet(doc, 'Operator MIN untuk implikasi (AND pada bagian IF)')
add_bullet(doc, 'Operator MAX untuk agregasi (menggabungkan beberapa aturan)')
add_bullet(doc, 'Defuzzifikasi menggunakan metode Centroid (Pusat Area)')

doc.add_heading('3.3 Empat Tahapan Proses Fuzzy Mamdani', level=2)

# --- Tahap 1: Fuzzifikasi ---
p = doc.add_paragraph()
run = p.add_run('Tahap 1: FUZZIFIKASI')
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Times New Roman'

p = doc.add_paragraph('Mengubah nilai tegas (crisp) menjadi derajat keanggotaan fuzzy (0.0 – 1.0). '
    'Pada tahap ini, angka pasti dari input dikonversi menjadi nilai keanggotaan '
    'menggunakan Fungsi Keanggotaan (Membership Function) berbentuk segitiga (triangular).')
p.runs[0].font.name = 'Times New Roman'

p = doc.add_paragraph()
run = p.add_run('Rumus Fungsi Segitiga:')
run.bold = True
run.font.name = 'Times New Roman'

add_formula(doc, 'μ(x) = 0,                    jika x ≤ a atau x ≥ c')
add_formula(doc, 'μ(x) = (x − a) / (b − a),    jika a < x < b')
add_formula(doc, 'μ(x) = (c − x) / (c − b),    jika b < x < c')
add_formula(doc, 'μ(x) = 1,                    jika x = b')

p = doc.add_paragraph('Dimana: a = batas kiri, b = titik puncak, c = batas kanan')
p.runs[0].font.name = 'Times New Roman'
p.runs[0].italic = True

# --- Tahap 2: Inferensi ---
p = doc.add_paragraph()
run = p.add_run('Tahap 2: INFERENSI (Evaluasi Aturan)')
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Times New Roman'

p = doc.add_paragraph('Mencocokkan nilai fuzzifikasi dengan basis aturan IF-THEN. '
    'Sistem menggunakan 243 aturan yang mengkombinasikan 5 variabel input.')
p.runs[0].font.name = 'Times New Roman'

p = doc.add_paragraph()
run = p.add_run('Format Aturan:')
run.bold = True
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.27)
run = p.add_run(
    'JIKA Budget ADALAH [Rendah/Sedang/Tinggi]\n'
    'DAN  Level PC ADALAH [Rendah/Sedang/Tinggi]\n'
    'DAN  Tipe Gamer ADALAH [Casual/Balanced/Hardcore]\n'
    'DAN  Rating ADALAH [Rendah/Sedang/Tinggi]\n'
    'DAN  Waktu Bermain ADALAH [Singkat/Sedang/Lama]\n'
    'MAKA Rekomendasi ADALAH [Tidak Direkomendasikan / Kurang Direkomendasikan / Direkomendasikan / Sangat Direkomendasikan]'
)
run.font.name = 'Consolas'
run.font.size = Pt(10)

p = doc.add_paragraph()
run = p.add_run('Operator yang digunakan: ')
run.font.name = 'Times New Roman'
run = p.add_run('Fungsi MIN (nilai minimum) untuk operator AND.')
run.bold = True
run.font.name = 'Times New Roman'

add_formula(doc, 'α-predikat = MIN(μ_budget, μ_pc_level, μ_gamer_type, μ_rating, μ_playtime)')

# --- Tahap 3: Agregasi ---
p = doc.add_paragraph()
run = p.add_run('Tahap 3: AGREGASI')
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Times New Roman'

p = doc.add_paragraph('Menggabungkan semua aturan yang terpicu menjadi satu area kurva. '
    'Jika beberapa aturan menghasilkan kategori output yang sama, sistem mengambil '
    'nilai MAX (nilai tertinggi) dari kekuatan aturan tersebut.')
p.runs[0].font.name = 'Times New Roman'

add_formula(doc, 'μ_output(kategori) = MAX(α₁, α₂, α₃, ..., αₙ)')

# --- Tahap 4: Defuzzifikasi ---
p = doc.add_paragraph()
run = p.add_run('Tahap 4: DEFUZZIFIKASI (Metode Centroid)')
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Times New Roman'

p = doc.add_paragraph('Mengubah area fuzzy menjadi satu nilai tegas (crisp) sebagai output akhir. '
    'Metode Centroid (Pusat Area / Center of Area) menghitung titik pusat massa '
    'dari area di bawah kurva gabungan.')
p.runs[0].font.name = 'Times New Roman'

p = doc.add_paragraph()
run = p.add_run('Rumus Centroid (kontinu):')
run.bold = True
run.font.name = 'Times New Roman'
add_formula(doc, 'z* = ∫ x · μ(x) dx  /  ∫ μ(x) dx')

p = doc.add_paragraph()
run = p.add_run('Implementasi diskrit (pada kode, 1000 titik sampel pada rentang 0–100):')
run.bold = True
run.font.name = 'Times New Roman'
add_formula(doc, 'z* = Σᵢ₌₀¹⁰⁰⁰ (xᵢ · μ(xᵢ))  /  Σᵢ₌₀¹⁰⁰⁰ μ(xᵢ)')

p = doc.add_paragraph('Hasil defuzzifikasi ini adalah Skor Rekomendasi (0 – 100).')
p.runs[0].font.name = 'Times New Roman'
p.runs[0].bold = True

doc.add_page_break()

# ============================================================
# BAB 4: CARA KERJA CODINGAN
# ============================================================
doc.add_heading('4. CARA KERJA CODINGAN (Arsitektur Sistem)', level=1)
add_line(doc)

doc.add_heading('4.1 Struktur File Fuzzy Engine', level=2)

add_table(doc,
    ['File', 'Fungsi'],
    [
        ['membership.py', 'Fungsi keanggotaan segitiga (kurva)'],
        ['fuzzification.py', 'Mengubah angka crisp → derajat keanggotaan'],
        ['inference.py', '243 aturan IF-THEN + evaluasi aturan'],
        ['aggregation.py', 'Menggabungkan aturan (operator MAX)'],
        ['defuzzification.py', 'Menghitung centroid → skor akhir (0-100)'],
        ['recommendation.py', 'Orkestrasi keseluruhan + koneksi database'],
    ],
    col_widths=[4.5, 11]
)

doc.add_heading('4.2 Alur Kode Program (Step-by-Step)', level=2)

steps = [
    ('1. PENGGUNA', 'Mengisi form (recommend.html) → Budget, PC Level, Tipe Gamer, Rating, Waktu Bermain, Genre'),
    ('2. FILTER DATABASE', 'Query MySQL: SELECT games WHERE genre LIKE \'%Action%\' → Membuang game yang genre-nya tidak sesuai'),
    ('3. FUZZIFIKASI', 'Setiap game dikonversi ke derajat keanggotaan (fuzzification.py + membership.py)'),
    ('4. INFERENSI', 'Evaluasi 243 aturan menggunakan operator MIN (inference.py)'),
    ('5. AGREGASI', 'Menggabungkan kekuatan aturan dengan operator MAX (aggregation.py)'),
    ('6. DEFUZZIFIKASI', 'Metode Centroid → Skor kecocokan (defuzzification.py)'),
    ('7. PERINGKAT & OUTPUT', 'Urutkan skor tertinggi → terendah → Ambil Top 10 → Tampilkan di browser'),
]

add_table(doc,
    ['Langkah', 'Proses'],
    [[s[0], s[1]] for s in steps],
    col_widths=[4.5, 11]
)

doc.add_heading('4.3 Variabel Fuzzy yang Digunakan', level=2)

p = doc.add_paragraph()
run = p.add_run('Variabel Input (5 variabel):')
run.bold = True
run.font.name = 'Times New Roman'

add_table(doc,
    ['No', 'Variabel', 'Himpunan Fuzzy', 'Rentang'],
    [
        ['1', 'Anggaran (Budget)', 'Rendah, Sedang, Tinggi', 'Rp 0 – Rp 1.000.000'],
        ['2', 'Level PC', 'Rendah, Sedang, Tinggi', '1, 2, 3 (singleton)'],
        ['3', 'Tipe Gamer', 'Casual, Balanced, Hardcore', '1, 2, 3 (singleton)'],
        ['4', 'Rating', 'Rendah, Sedang, Tinggi', '0% – 100%'],
        ['5', 'Waktu Bermain', 'Singkat, Sedang, Lama', '0 – 200 jam'],
    ],
    col_widths=[1.2, 4, 5, 5]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Variabel Output (1 variabel):')
run.bold = True
run.font.name = 'Times New Roman'

add_table(doc,
    ['Kategori Output', 'Rentang Skor'],
    [
        ['Tidak Direkomendasikan', '0 – 25'],
        ['Kurang Direkomendasikan', '20 – 50'],
        ['Direkomendasikan', '45 – 75'],
        ['Sangat Direkomendasikan', '70 – 100'],
    ],
    col_widths=[7, 5]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Total Aturan Fuzzy:')
run.bold = True
run.font.name = 'Times New Roman'
add_formula(doc, '3(budget) × 3(pc_level) × 3(gamer_type) × 3(rating) × 3(playtime) = 243 Aturan')

doc.add_page_break()

# ============================================================
# BAB 5: CONTOH INPUT → PROSES → OUTPUT
# ============================================================
doc.add_heading('5. CONTOH INPUT → PROSES → OUTPUT', level=1)
add_line(doc)

doc.add_heading('5.1 Skenario: Mahasiswa Mencari Game Action', level=2)

p = doc.add_paragraph()
run = p.add_run('📥 INPUT (Pengguna memasukkan preferensi di form web)')
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Times New Roman'

add_table(doc,
    ['Parameter', 'Nilai yang Dimasukkan'],
    [
        ['Anggaran', 'Rp 300.000'],
        ['Spesifikasi PC', 'Medium (Mid-End)'],
        ['Tipe Gamer', 'Balanced (Seimbang)'],
        ['Rating Minimum', '75%'],
        ['Waktu Bermain', '30 Jam'],
        ['Genre', 'Action'],
    ],
    col_widths=[5, 8]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('⚙️ PROSES (Apa yang terjadi di dalam sistem)')
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Times New Roman'

# Langkah 1
p = doc.add_paragraph()
run = p.add_run('Langkah 1: Filter Database')
run.bold = True
run.underline = True
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.27)
run = p.add_run(
    "SELECT * FROM games\n"
    "WHERE (genre LIKE '%Action%' OR tags LIKE '%Action%')\n"
    "  AND total_reviews >= 5\n"
    "  AND price_idr <= 300000"
)
run.font.name = 'Consolas'
run.font.size = Pt(10)

p = doc.add_paragraph()
run = p.add_run('Hasil: ')
run.bold = True
run.font.name = 'Times New Roman'
run = p.add_run('Sistem menemukan sejumlah game Action yang sesuai budget dari database.')
run.font.name = 'Times New Roman'

# Langkah 2: Fuzzifikasi
p = doc.add_paragraph()
run = p.add_run('Langkah 2: Fuzzifikasi (untuk game peringkat #1)')
run.bold = True
run.underline = True
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
run = p.add_run('Contoh Game Kandidat: "Nuclear Option" (Harga: Rp 287.840, Rating: 94%, 7.024 ulasan)')
run.bold = True
run.font.name = 'Times New Roman'

# Budget
p = doc.add_paragraph()
run = p.add_run('A. Fuzzifikasi Budget (Harga game Rp 287.840):')
run.bold = True
run.font.name = 'Times New Roman'

add_formula(doc, 'μ_Low(287840) = (c − x) / (c − b) = (300000 − 287840) / (300000 − 0) = 0.04')
add_formula(doc, 'μ_Medium(287840) = (x − a) / (b − a) = (287840 − 50000) / (300000 − 50000) = 0.95')
add_formula(doc, 'μ_High(287840) = 0.00   (di luar rentang, x < a = 500000)')

p = doc.add_paragraph()
run = p.add_run('Hasil Fuzzifikasi Budget:')
run.bold = True
run.font.name = 'Times New Roman'
add_formula(doc, 'Budget = { Low: 0.04,  Medium: 0.95,  High: 0.00 }')

# PC Level
p = doc.add_paragraph()
run = p.add_run('B. Fuzzifikasi Level PC (Game butuh level 3/High, User punya level 2/Medium):')
run.bold = True
run.font.name = 'Times New Roman'
add_formula(doc, 'PC Level = { Low: 0.0,  Medium: 1.0,  High: 0.0 }  (Singleton: nilai PC user = 2)')

# Gamer Type
p = doc.add_paragraph()
run = p.add_run('C. Fuzzifikasi Tipe Gamer (User memilih Balanced = 2):')
run.bold = True
run.font.name = 'Times New Roman'
add_formula(doc, 'Gamer Type = { Casual: 0.0,  Balanced: 1.0,  Hardcore: 0.0 }')

# Rating
p = doc.add_paragraph()
run = p.add_run('D. Fuzzifikasi Rating (Rating game 94%, preferensi user 75%):')
run.bold = True
run.font.name = 'Times New Roman'
add_formula(doc, 'μ_High(94) = (x − a) / (b − a) = (94 − 60) / (100 − 60) = 0.85')
add_formula(doc, 'Rating = { Low: 0.0,  Medium: 0.0,  High: 0.85 }')

# Playtime
p = doc.add_paragraph()
run = p.add_run('E. Fuzzifikasi Waktu Bermain (Game ~30 jam, preferensi user 30 jam):')
run.bold = True
run.font.name = 'Times New Roman'
add_formula(doc, 'Playtime = { Short: 0.0,  Medium: 1.0,  Long: 0.0 }  (kecocokan sempurna)')

# Langkah 3: Inferensi
p = doc.add_paragraph()
run = p.add_run('Langkah 3: Inferensi (Evaluasi Aturan)')
run.bold = True
run.underline = True
run.font.name = 'Times New Roman'

p = doc.add_paragraph('Dari 243 aturan, yang paling relevan untuk game ini misalnya:')
p.runs[0].font.name = 'Times New Roman'

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.27)
run = p.add_run(
    'JIKA Budget = Medium (0.95)\n'
    'DAN  PC Level = Medium (1.0)\n'
    'DAN  Gamer Type = Balanced (1.0)\n'
    'DAN  Rating = High (0.85)\n'
    'DAN  Playtime = Medium (1.0)\n'
    'MAKA Rekomendasi = Sangat Direkomendasikan'
)
run.font.name = 'Consolas'
run.font.size = Pt(10)

p = doc.add_paragraph()
run = p.add_run('Menghitung α-predikat (menggunakan operator MIN):')
run.bold = True
run.font.name = 'Times New Roman'

add_formula(doc, 'α = MIN(0.95, 1.0, 1.0, 0.85, 1.0) = 0.85')

p = doc.add_paragraph('Profil ideal untuk tipe Balanced: {budget: Medium, rating: High, playtime: Medium}')
p.runs[0].font.name = 'Times New Roman'
add_bullet(doc, 'Budget game = Medium, ideal Balanced = Medium → cocok ✓')
add_bullet(doc, 'Rating game = High, ideal Balanced = High → cocok ✓')
add_bullet(doc, 'Playtime game = Medium, ideal Balanced = Medium → cocok ✓')

p = doc.add_paragraph()
run = p.add_run('Kecocokan = 3 dari 3 → consequent = ')
run.font.name = 'Times New Roman'
run = p.add_run('Sangat Direkomendasikan')
run.bold = True
run.font.name = 'Times New Roman'
run = p.add_run(' (index 3). Kekuatan aturan untuk kategori "Sangat Direkomendasikan" = 0.85')
run.font.name = 'Times New Roman'

# Langkah 4: Agregasi
p = doc.add_paragraph()
run = p.add_run('Langkah 4: Agregasi')
run.bold = True
run.underline = True
run.font.name = 'Times New Roman'

p = doc.add_paragraph('Setelah semua 243 aturan dievaluasi, hasil agregasi (MAX dari setiap kategori):')
p.runs[0].font.name = 'Times New Roman'

add_table(doc,
    ['Kategori', 'Strength (Kekuatan)'],
    [
        ['Not Recommended', '0.00'],
        ['Less Recommended', '0.04'],
        ['Recommended', '0.85'],
        ['Highly Recommended', '0.85 ← Nilai tertinggi'],
    ],
    col_widths=[6, 6]
)

# Langkah 5: Defuzzifikasi
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Langkah 5: Defuzzifikasi (Centroid)')
run.bold = True
run.underline = True
run.font.name = 'Times New Roman'

p = doc.add_paragraph('Sistem menghitung titik pusat massa (centroid) dari area di bawah kurva:')
p.runs[0].font.name = 'Times New Roman'

add_formula(doc, 'μ(xᵢ) = MAX_kategori ( MIN( μ_mf_kategori(xᵢ), α_kategori ) )')
add_formula(doc, 'z* = Σᵢ₌₀¹⁰⁰⁰ (xᵢ · μ(xᵢ))  /  Σᵢ₌₀¹⁰⁰⁰ μ(xᵢ)')

p = doc.add_paragraph()
run = p.add_run('Hasil Defuzzifikasi: Skor = 79.0')
run.bold = True
run.font.name = 'Times New Roman'

# Langkah 6
p = doc.add_paragraph()
run = p.add_run('Langkah 6: Kategorisasi Skor')
run.bold = True
run.underline = True
run.font.name = 'Times New Roman'

add_formula(doc, 'z* = 79.0  →  70 ≤ 79.0 ≤ 100  →  Kategori: Sangat Direkomendasikan ✓')

# OUTPUT
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('📤 OUTPUT (Ditampilkan di browser pengguna)')
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Times New Roman'

p = doc.add_paragraph('Setelah semua game Action dievaluasi satu per satu, sistem mengurutkan skor tertinggi dan menampilkan Top 10:')
p.runs[0].font.name = 'Times New Roman'

add_table(doc,
    ['#', 'Nama Game', 'Harga', 'Rating', 'Ulasan', 'Skor', 'Kategori'],
    [
        ['1', 'Nuclear Option', 'Rp 287.840', '94%', '7.024', '79%', 'Sangat Direkomendasikan'],
        ['2', 'Dome Keeper', 'Rp 287.840', '91%', '16.108', '77%', 'Sangat Direkomendasikan'],
        ['3', 'Life is Strange: True Colors', 'Rp 287.840', '88%', '12.015', '75%', 'Direkomendasikan'],
        ['4', 'Over The Top: WWI', 'Rp 260.160', '88%', '3.572', '72%', 'Direkomendasikan'],
        ['5', 'Castle Crashers®', 'Rp 239.840', '95%', '105.576', '72%', 'Direkomendasikan'],
        ['6', 'DRAGON BALL FighterZ', 'Rp 230.240', '91%', '53.794', '69%', 'Direkomendasikan'],
        ['7', 'People Playground', 'Rp 159.840', '97%', '288.135', '69%', 'Direkomendasikan'],
        ['8', 'Crime Simulator', 'Rp 239.840', '85%', '4.499', '68%', 'Direkomendasikan'],
        ['9', 'End of Zoe', 'Rp 239.840', '85%', '1.190', '68%', 'Direkomendasikan'],
        ['10', 'LORT', 'Rp 239.840', '83%', '2.759', '68%', 'Direkomendasikan'],
    ],
    col_widths=[1, 4, 2.5, 1.5, 2, 1.5, 4]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Alasan yang ditampilkan untuk setiap game (Explainable AI):')
run.bold = True
run.font.name = 'Times New Roman'
add_bullet(doc, '✅ "Cocok untuk gaya bermain Seimbang"')
add_bullet(doc, '✅ "Rp 287.840 sesuai dengan anggaran Rp 300.000 Anda"')
add_bullet(doc, '✅ "Rating 94% memenuhi preferensi 75% Anda"')
add_bullet(doc, '⚠️ "Game ini membutuhkan spesifikasi High End (mungkin berat untuk PC Anda)"')

doc.add_page_break()

# ============================================================
# BAB 6: REFERENSI
# ============================================================
doc.add_heading('6. REFERENSI JURNAL', level=1)
add_line(doc)

doc.add_heading('6.1 Referensi Utama (Pencetus Metode)', level=2)

refs_main = [
    '[1]  E. H. Mamdani and S. Assilian, "An experiment in linguistic synthesis with a fuzzy logic controller," '
    'International Journal of Man-Machine Studies, Vol. 7, No. 1, pp. 1-13, 1975.',
    '[2]  L. A. Zadeh, "Fuzzy Sets," Information and Control, Vol. 8, No. 3, pp. 338-353, 1965.',
]

for ref in refs_main:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

doc.add_heading('6.2 Referensi Pendukung', level=2)

refs_support = [
    '[3]  S. Kusumadewi and H. Purnomo, "Aplikasi Logika Fuzzy untuk Pendukung Keputusan," '
    'Graha Ilmu, Yogyakarta, 2010.',
    '[4]  T. J. Ross, "Fuzzy Logic with Engineering Applications," John Wiley & Sons, 3rd Edition, 2010.',
    '[5]  F. Ricci, L. Rokach, B. Shapira, and P. B. Kantor, "Recommender Systems Handbook," Springer, 2011.',
    '[6]  J. M. Mendel, "Fuzzy Logic Systems for Engineering: A Tutorial," Proceedings of the IEEE, '
    'Vol. 83, No. 3, pp. 345-377, 1995.',
]

for ref in refs_support:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

doc.add_page_break()

# ============================================================
# BAB 7: BATASAN
# ============================================================
doc.add_heading('7. BATASAN / KEKURANGAN SISTEM', level=1)
add_line(doc)

add_table(doc,
    ['No', 'Kekurangan', 'Penjelasan'],
    [
        ['1', 'Data Statis', 'Database game diimpor secara manual (tidak real-time dari Steam API)'],
        ['2', 'Genre Bukan Variabel Fuzzy', 'Genre hanya digunakan sebagai filter kasar, tidak dievaluasi dalam logika fuzzy'],
        ['3', 'Tidak Ada Personalisasi', 'Sistem tidak menyimpan riwayat pencarian pengguna (tanpa akun/login)'],
        ['4', 'Spek PC Diestimasi', 'Level PC game ditentukan dari tags/genre (bukan dari data system requirements resmi)'],
        ['5', 'Jumlah Aturan Besar', '243 aturan sulit dikelola manual jika variabel bertambah (skalabilitas terbatas)'],
    ],
    col_widths=[1.2, 4.5, 10]
)

doc.add_page_break()

# ============================================================
# BAB 8: KESIMPULAN
# ============================================================
doc.add_heading('8. RINGKASAN UNTUK PENUTUP PRESENTASI', level=1)
add_line(doc)

doc.add_heading('Kesimpulan', level=2)

conclusions = [
    'GameFinder berhasil mengimplementasikan metode Fuzzy Mamdani sebagai mesin inferensi untuk sistem rekomendasi game Steam.',
    'Sistem menggunakan 5 variabel input dan 243 aturan fuzzy untuk mengevaluasi kecocokan setiap game dengan profil pengguna.',
    'Proses Fuzzifikasi → Inferensi → Agregasi → Defuzzifikasi (Centroid) menghasilkan skor rekomendasi 0-100 yang dapat dijelaskan (Explainable AI).',
    'Sistem menampilkan Top 10 game terbaik beserta alasan rekomendasinya.',
]

for i, c in enumerate(conclusions, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. ')
    run.bold = True
    run.font.name = 'Times New Roman'
    run = p.add_run(c)
    run.font.name = 'Times New Roman'

doc.add_heading('Saran Pengembangan', level=2)

suggestions = [
    'Integrasi dengan Steam API untuk data game real-time.',
    'Penambahan fitur akun pengguna untuk personalisasi riwayat.',
    'Penggunaan metode hybrid (Fuzzy + Machine Learning) untuk akurasi lebih tinggi.',
    'Penambahan variabel fuzzy baru seperti popularitas dan umur game.',
]

for i, s in enumerate(suggestions, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. ')
    run.bold = True
    run.font.name = 'Times New Roman'
    run = p.add_run(s)
    run.font.name = 'Times New Roman'

# ============================================================
# SAVE
# ============================================================
output_path = r'd:\TUGAS KULIAH GHALIB\2026 GENAP SMS 4\Kecerdasan Buatan\GameFinder\PRESENTASI_GAMEFINDER_v3.docx'
doc.save(output_path)
print(f'File berhasil disimpan: {output_path}')
